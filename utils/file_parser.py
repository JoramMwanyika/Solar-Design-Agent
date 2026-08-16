"""
File parser — extracts text content from uploaded files.
Supports PDF, DOCX, Excel, plain text, and passes images to Gemini Vision.
"""
import io
from pathlib import Path
import warnings


def parse_uploaded_file(file_bytes: bytes, filename: str) -> str:
    """
    Parses an uploaded file and returns its text content.
    Returns empty string for image files (handled separately by Gemini Vision).
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_bytes)
    elif ext in (".xlsx", ".xls"):
        return _parse_excel(file_bytes)
    elif ext in (".txt", ".csv"):
        return file_bytes.decode("utf-8", errors="ignore")
    elif ext in (".png", ".jpg", ".jpeg", ".webp", ".gif"):
        return ""  # Images handled by Gemini Vision
    else:
        # Try UTF-8 decode as fallback
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""


def _parse_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            # Limit to first 5 pages for token efficiency and performance
            for page in pdf.pages[:5]:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
                # Also extract tables
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        text_parts.append(" | ".join(str(c or "") for c in row))
        return "\n".join(text_parts)
    except ImportError:
        return ""
    except Exception as e:
        return f"[PDF parse error: {e}]"


def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)
    except ImportError:
        return ""
    except Exception as e:
        return f"[DOCX parse error: {e}]"


def _parse_excel(file_bytes: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(c) if c is not None else "" for c in row)
                if row_text.strip(" |"):
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        return ""
    except Exception as e:
        return f"[Excel parse error: {e}]"


def get_mime_type(filename: str) -> str:
    """Returns MIME type from filename extension."""
    ext = Path(filename).suffix.lower()
    mime_map = {
        ".pdf": "application/pdf",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".txt": "text/plain",
        ".csv": "text/csv",
        ".json": "application/json",
    }
    return mime_map.get(ext, "application/octet-stream")


def extract_loads_from_dataframe(df_raw) -> tuple[list[dict], bool]:
    """
    Intelligently extracts load items from any uploaded CSV/Excel DataFrame.
    Automatically handles both:
      1) Standard Appliance Schedules / Load Profiles (summing items & using explicit energy/time columns)
      2) Logged Meter Time-Series Data (peak interval & daily averages)
    Returns: (loads_list, is_time_series_flag)
    """
    import pandas as pd
    import re

    if df_raw is None or df_raw.empty:
        return [], False

    df = df_raw.copy()

    # Step 1: Check if top rows are metadata/header rows
    def looks_like_header(row_vals):
        words = set()
        for v in row_vals:
            if pd.notna(v):
                for w in str(v).lower().replace("(", " ").replace(")", " ").replace("[", " ").replace("]", " ").replace("-", " ").replace("_", " ").split():
                    words.add(w)
        power_hits = sum(1 for kw in ("watt", "watts", "wattage", "active", "apparent", "power", "kw", "kva", "energy") if kw in words)
        if power_hits >= 1:
            return True
        general_hits = sum(1 for kw in ("loads", "name", "appliance", "item", "description", "timestamp", "time", "date", "interval", "load", "demand", "quantity", "qty", "hours") if kw in words)
        return general_hits >= 2

    if not looks_like_header(df.columns):
        for idx in range(min(20, len(df))):
            row_vals = df.iloc[idx].values
            if looks_like_header(row_vals):
                df.columns = [str(v).strip() if pd.notna(v) else f"col_{i}" for i, v in enumerate(row_vals)]
                df = df.iloc[idx + 1:].reset_index(drop=True)
                break

    # Clean column names for mapping
    col_map = {str(c).lower().strip(): c for c in df.columns}

    def find_col(*candidates):
        # 1. Try exact matches first
        for cand in candidates:
            if cand in col_map:
                return col_map[cand]
        # 2. Try substring matches
        for cand in candidates:
            for c_low, c_orig in col_map.items():
                if cand in c_low and not c_low.startswith("unnamed"):
                    return c_orig
        return None

    # Candidate columns classification
    col_energy = find_col("energy(kwh)", "energy (kwh)", "daily energy", "total energy", "kwh/day", "energy", "kwh", "energy(wh)", "energy (wh)", "wh/day", "wh")
    col_total_p = find_col("total power", "total_power", "total kw", "total_kw", "total w", "total p")
    col_hrs = find_col("time (hrs)", "time(hrs)", "hours_per_day", "hours", "hrs", "duration", "runtime", "operating hours", "h/day")
    col_qty = find_col("quantity", "qty", "count", "num", "units", "pcs")
    col_name = find_col("loads", "load", "name", "appliance", "item", "description", "device", "equipment", "timestamp", "date & time", "date_time", "datetime")
    col_unit_p = find_col("power(kw)", "power (kw)", "power", "wattage", "active power", "active_power", "active", "kw", "watt", "watts", "w")
    col_s = find_col("apparent wattage", "apparent_wattage", "apparent power", "apparent_power", "apparent", "s (kva)", "s (va)", "kva", "va")

    col_diversity = find_col("diversity factor", "diversity_factor", "diversity", "df")
    col_utilisation = find_col("utilisation factor", "utilization factor", "utilisation_factor", "utilization_factor", "utilisation", "utilization", "uf")

    if col_unit_p == col_s and col_s is not None:
        col_unit_p = None
    if col_total_p == col_s and col_s is not None:
        col_total_p = None

    col_date_only = find_col("date", "day")
    col_time_only = find_col("time of day", "timestamp", "date & time", "datetime", "date_time", "time", "clock")

    # Time-series log detection & Date/Time Combination
    is_time_series = False
    if col_name and any(kw in str(col_name).lower() for kw in ("timestamp", "interval", "date")):
        is_time_series = True
    elif len(df) >= 20 and not col_qty and not col_hrs and not col_energy:
        is_time_series = True

    datetime_col_name = None
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        if col_date_only and col_time_only and col_date_only != col_time_only:
            try:
                combined_series = df[col_date_only].astype(str) + " " + df[col_time_only].astype(str)
                parsed_dt = pd.to_datetime(combined_series, errors="coerce")
                if parsed_dt.notna().sum() > 0:
                    df["__parsed_dt__"] = parsed_dt
                    datetime_col_name = "__parsed_dt__"
                    is_time_series = True
            except Exception:
                pass

        if datetime_col_name is None:
            for c in (col_name, col_time_only, col_date_only):
                if c is not None:
                    try:
                        parsed_dt = pd.to_datetime(df[c], errors="coerce")
                        if parsed_dt.notna().sum() >= max(3, len(df) * 0.3):
                            df["__parsed_dt__"] = parsed_dt
                            datetime_col_name = "__parsed_dt__"
                            is_time_series = True
                            break
                    except Exception:
                        pass

    # Sort data chronologically if date/time column is detected
    if datetime_col_name is not None and df[datetime_col_name].notna().sum() > 0:
        df = df.sort_values(by=datetime_col_name, na_position="last").reset_index(drop=True)

    def get_val(r, c):
        if c is None:
            return None
        v = r.get(c)
        if isinstance(v, pd.Series):
            return v.iloc[0]
        return v

    def parse_number(val):
        if pd.isna(val):
            return None
        if isinstance(val, (int, float)):
            return float(val) if val == val else None
        s = str(val).replace(",", "").replace("$", "").strip()
        m = re.search(r"[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?", s)
        if m:
            return float(m.group(0))
        return None

    csv_loads = []
    for idx, row in df.iterrows():
        c_name_val = get_val(row, col_name)
        c_name_str = str(c_name_val or "").strip()
        c_name_lower = c_name_str.lower()

        # Skip summary rows and section header rows
        if any(kw in c_name_lower for kw in ("daily energy demand", "total energy", "subtotal", "grand total", "total demand")):
            continue

        qty_val = parse_number(get_val(row, col_qty)) if col_qty else 1.0
        qty = int(qty_val) if (qty_val is not None and qty_val > 0) else 1

        hrs_val = parse_number(get_val(row, col_hrs)) if col_hrs else None
        energy_val = parse_number(get_val(row, col_energy)) if col_energy else None

        div_val = parse_number(get_val(row, col_diversity)) if col_diversity else 1.0
        diversity = float(div_val) if (div_val is not None and div_val > 0) else 1.0

        util_val = parse_number(get_val(row, col_utilisation)) if col_utilisation else 1.0
        utilisation = float(util_val) if (util_val is not None and util_val > 0) else 1.0

        unit_p_w = 0.0
        apparent_va = None

        # 1. Total Power Column Processing
        if col_total_p is not None:
            raw_tot_val = get_val(row, col_total_p)
            raw_tot = parse_number(raw_tot_val)
            if raw_tot is not None and raw_tot > 0:
                raw_tot_str = str(raw_tot_val or "").lower()
                col_lbl = str(col_total_p).lower()
                if "kw" in raw_tot_str or "kw" in col_lbl:
                    is_kw = True
                elif "w" in raw_tot_str or "watt" in raw_tot_str or "w" in col_lbl:
                    is_kw = False
                else:
                    is_kw = raw_tot < 2.0
                unit_p_w = (raw_tot * (1000.0 if is_kw else 1.0)) / qty

        # 2. Unit Power Column Processing
        elif col_unit_p is not None:
            raw_p_val = get_val(row, col_unit_p)
            raw_p = parse_number(raw_p_val)
            if raw_p is not None and raw_p > 0:
                raw_p_str = str(raw_p_val or "").lower()
                col_label = str(col_unit_p).lower()
                if "kw" in raw_p_str or "kw" in col_label:
                    is_kw = True
                elif "w" in raw_p_str or "watt" in raw_p_str or "w" in col_label:
                    is_kw = False
                else:
                    is_kw = raw_p < 2.0
                unit_p_w = raw_p * (1000.0 if is_kw else 1.0)

        # Read apparent power (kVA/VA) if available
        if col_s is not None:
            raw_s_val = get_val(row, col_s)
            raw_s = parse_number(raw_s_val)
            if raw_s is not None and raw_s > 0:
                raw_s_str = str(raw_s_val or "").lower()
                is_kva = "kva" in raw_s_str or "kva" in str(col_s).lower() or (is_time_series and raw_s < 500.0)
                apparent_va = raw_s * (1000.0 if is_kva else 1.0)

        # Fallback: if no active power column but apparent power exists, derive active power
        if unit_p_w == 0.0 and apparent_va is not None and apparent_va > 0:
            unit_p_w = apparent_va * 0.85  # assume PF = 0.85

        # Helper to get energy in Wh based on column unit or magnitude
        def get_energy_in_wh(e_val, col_e, cell_raw_val):
            if e_val is None or e_val <= 0:
                return None
            cell_str = str(cell_raw_val or "").lower()
            col_lbl = str(col_e).lower() if col_e else ""
            if "kwh" in cell_str or "kwh" in col_lbl:
                return e_val * 1000.0
            elif "wh" in cell_str or "wh" in col_lbl:
                return e_val
            elif e_val >= 500.0:
                return e_val
            else:
                return e_val * 1000.0

        raw_energy_cell = get_val(row, col_energy)
        energy_wh_val = get_energy_in_wh(energy_val, col_energy, raw_energy_cell)

        # Fallback to energy / hrs calculation if unit_p_w == 0
        if unit_p_w == 0.0 and energy_wh_val is not None and energy_wh_val > 0 and hrs_val is not None and hrs_val > 0:
            unit_p_w = energy_wh_val / (qty * hrs_val * diversity * utilisation)

        # Determine operating hours
        if hrs_val is None or hrs_val <= 0:
            if energy_wh_val is not None and energy_wh_val > 0 and unit_p_w > 0:
                hrs_val = energy_wh_val / (unit_p_w * qty * diversity * utilisation)
            else:
                hrs_val = 12.0 if not is_time_series else 1.0

        # Calculate exact energy taking Diversity Factor and Utilisation Factor into account
        if energy_wh_val is None and unit_p_w > 0 and hrs_val > 0:
            energy_wh_val = unit_p_w * qty * hrs_val * diversity * utilisation

        if unit_p_w > 0 or (energy_wh_val is not None and energy_wh_val > 0):
            c_name_val = get_val(row, col_name)
            dt_val = get_val(row, "__parsed_dt__") if datetime_col_name else None
            
            if pd.notna(dt_val):
                try:
                    item_name = pd.to_datetime(dt_val).isoformat()
                except Exception:
                    item_name = str(c_name_val).strip() if (col_name and pd.notna(c_name_val) and str(c_name_val).strip()) else f"Load Item {idx+1}"
            else:
                item_name = str(c_name_val).strip() if (col_name and pd.notna(c_name_val) and str(c_name_val).strip()) else f"Load Item {idx+1}"

            # Ignore pure header rows (e.g. "1 Manager's House" with no rating, power, energy or hrs)
            if hrs_val <= 0 and energy_wh_val is None:
                continue

            explicit_wh = energy_wh_val

            csv_loads.append({
                "name": item_name,
                "wattage": unit_p_w * diversity * utilisation,  # Effective active wattage per unit
                "quantity": qty,
                "hours_per_day": float(hrs_val),
                "apparent_wattage": apparent_va,
                "is_time_series": is_time_series,
                "explicit_daily_energy_wh": explicit_wh
            })

    return csv_loads, is_time_series
